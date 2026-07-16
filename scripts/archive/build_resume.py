#!/usr/bin/env python3
"""Generates static/adam_robinson.docx from structured content."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

TEAL = RGBColor(0x2A, 0x7A, 0x7A)
BLACK = RGBColor(0x0A, 0x0A, 0x0A)
MUTED = RGBColor(0x44, 0x44, 0x44)
SITE_URL = 'https://www.adamrobinson.tech'
GITHUB_URL = 'https://github.com/aaddaamm'
LINKEDIN_URL = 'https://www.linkedin.com/in/adam-robinson-software/'
EMAIL = 'adam@adamrobinson.tech'


def clear_paragraph_spacing(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def add_rule(doc):
    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2A7A7A')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), relationship_id)

    run = OxmlElement('w:r')
    properties = OxmlElement('w:rPr')
    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:ascii'), 'Calibri')
    fonts.set(qn('w:hAnsi'), 'Calibri')
    size = OxmlElement('w:sz')
    size.set(qn('w:val'), '18')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '444444')
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'none')
    properties.extend([fonts, size, color, underline])
    run.append(properties)

    node = OxmlElement('w:t')
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_name_block(doc):
    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('adam robinson')
    run.font.name = 'JetBrains Mono'
    run.font.size = Pt(26)
    run.font.bold = False
    run.font.color.rgb = BLACK

    p2 = doc.add_paragraph()
    clear_paragraph_spacing(p2)
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run('SENIOR SOFTWARE ENGINEER')
    run2.font.name = 'JetBrains Mono'
    run2.font.size = Pt(9)
    run2.font.color.rgb = TEAL

    p3 = doc.add_paragraph()
    clear_paragraph_spacing(p3)
    p3.paragraph_format.space_after = Pt(2)
    run3 = p3.add_run('Providence, RI')
    run3.font.name = 'Calibri'
    run3.font.size = Pt(10)
    run3.font.color.rgb = MUTED

    p4 = doc.add_paragraph()
    clear_paragraph_spacing(p4)
    p4.paragraph_format.space_after = Pt(8)
    add_hyperlink(p4, 'adamrobinson.tech', SITE_URL)
    for label, url in [
        ('GitHub', GITHUB_URL),
        ('LinkedIn', LINKEDIN_URL),
        (EMAIL, f'mailto:{EMAIL}')
    ]:
        separator = p4.add_run('  ·  ')
        separator.font.name = 'Calibri'
        separator.font.size = Pt(9)
        separator.font.color.rgb = MUTED
        add_hyperlink(p4, label, url)


def add_section_heading(doc, text):
    add_rule(doc)
    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.font.name = 'JetBrains Mono'
    run.font.size = Pt(8)
    run.font.color.rgb = TEAL


def add_body(doc, text, space_after=4):
    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    return p


def add_role_header(doc, company, title, period, location=None):
    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)

    r1 = p.add_run(company)
    r1.font.name = 'Calibri'
    r1.font.size = Pt(10)
    r1.font.bold = True
    r1.font.color.rgb = BLACK

    r2 = p.add_run(f'  ·  {title}')
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10)
    r2.font.bold = False
    r2.font.color.rgb = MUTED

    p2 = doc.add_paragraph()
    clear_paragraph_spacing(p2)
    p2.paragraph_format.space_after = Pt(4)
    loc_text = f'{period}' + (f'  ·  {location}' if location else '')
    r3 = p2.add_run(loc_text)
    r3.font.name = 'JetBrains Mono'
    r3.font.size = Pt(8)
    r3.font.color.rgb = MUTED


def add_client_entry(doc, company, title, period, desc):
    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.18)

    r1 = p.add_run('-  ')
    r1.font.name = 'Calibri'
    r1.font.size = Pt(9)
    r1.font.color.rgb = TEAL

    r2 = p.add_run(company)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(9.5)
    r2.font.bold = True
    r2.font.color.rgb = BLACK

    r3 = p.add_run(f'  ·  {title}  ·  {period}')
    r3.font.name = 'Calibri'
    r3.font.size = Pt(9)
    r3.font.bold = False
    r3.font.color.rgb = MUTED

    p2 = doc.add_paragraph()
    clear_paragraph_spacing(p2)
    p2.paragraph_format.space_after = Pt(3)
    p2.paragraph_format.left_indent = Inches(0.35)
    r4 = p2.add_run(desc)
    r4.font.name = 'Calibri'
    r4.font.size = Pt(9.5)
    r4.font.color.rgb = BLACK


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    clear_paragraph_spacing(p)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK


def add_skill_row(doc, category, items):
    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    p.paragraph_format.space_after = Pt(2)

    r1 = p.add_run(f'{category:<12}')
    r1.font.name = 'JetBrains Mono'
    r1.font.size = Pt(9)
    r1.font.color.rgb = TEAL

    r2 = p.add_run(items)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10)
    r2.font.color.rgb = MUTED


def build():
    doc = Document()
    doc.core_properties.title = 'Adam Robinson - Senior Software Engineer'
    doc.core_properties.author = 'Adam Robinson'
    doc.core_properties.subject = 'Resume for Adam Robinson, Senior Software Engineer and Technical Lead'
    doc.core_properties.keywords = (
        'senior software engineer, technical lead, React, TypeScript, Node.js, Ruby on Rails, '
        'embedded consultant'
    )

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Default paragraph spacing
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    add_name_block(doc)

    # Summary
    add_section_heading(doc, 'Summary')
    add_body(doc,
        'Senior software engineer and technical lead with 15+ years delivering production systems '
        'across 15+ client engagements. Full-stack work in React, TypeScript/Node.js, and Rails '
        'across fintech, healthcare, industrial technology, and enterprise products. Recent work '
        'includes Rails bulk processing for nominee investments and a Strapi/React publishing '
        'pipeline with unified Auth0 login.'
    )

    # Experience
    add_section_heading(doc, 'Experience')

    add_role_header(doc, 'MojoTech', 'Senior Software Engineer / Technical Lead', 'Feb 2015 - Present', 'Providence, RI')
    add_body(doc,
        'Delivered 15+ client projects as an embedded senior engineer and technical lead. '
        'Selected clients:',
        space_after=6
    )

    clients = [
        ('iCapital', 'Senior Software Engineer, Consultant', 'May 2024 - present',
         'Co-designed a Rails service that consolidated bulk nominee processing for thousands '
         'of investments. Expanded localization across static and database-backed content and '
         "led the team's Supernova v1-to-v2 component library migration."),
        ('Healthcasts', 'Technical Lead', 'Oct 2022 - May 2024',
         'Led phased modernization of a medical publishing platform. Built a Strapi/React publishing '
         'pipeline and updated AWS infrastructure and frameworks. Unified authentication across '
         'products with Auth0, reducing publishing friction and unblocking an AI initiative.'),
        ('Angi', 'Senior Software Engineer, Consultant', 'Nov 2020 - Sep 2022',
         "Shipped across three post-merger codebases for HomeAdvisor, Handy, and Angie's List "
         'using Vue/Java, Rails/React, and Next.js/Contentful. Mentored interns through their first '
         'production release, a Careers page revamp.'),
        ('Shell Techworks', 'Software Engineer', 'Jun 2018 - Jul 2019',
         'Built a React and Node.js application that evaluated least-cost decommissioning paths '
         'for end-of-life offshore oil platforms. Used an onsite Google Design Sprint to narrow '
         'scope and deliver the MVP on schedule.'),
    ]
    for company, title, period, desc in clients:
        add_client_entry(doc, company, title, period, desc)

    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.35)
    r = p.add_run('Earlier clients include School of Motion, Amica Mutual, and AutoRaptor.')
    r.font.name = 'Calibri'
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = MUTED

    add_role_header(doc, 'Beacon Mutual Insurance', 'Associate Developer / Production Control', 'Mar 2011 - Feb 2015', 'Warwick, RI')
    add_body(doc,
        'Progressed from deployment and environment management to full-stack and database development, '
        'building claims, policy, financial transaction, and payment systems in a regulated environment.'
    )

    # Skills
    add_section_heading(doc, 'Skills')
    add_skill_row(doc, 'Backend', 'TypeScript, Ruby, SQL, Elixir  ·  Node.js, Ruby on Rails, Express, Phoenix')
    add_skill_row(doc, 'Frontend', 'React, SvelteKit, Vue, Next.js')
    add_skill_row(doc, 'Platform', 'AWS, Vercel, GitHub Actions  ·  Git, Prisma, Strapi, Contentful, Auth0')
    add_skill_row(doc, 'AI', 'Codex, Claude, Pi, GitHub Copilot  ·  agent instructions, skills, verification workflows')

    out = '/Users/adam/dotcom/static/adam_robinson.docx'
    doc.save(out)
    print(f'Saved: {out}')


if __name__ == '__main__':
    build()
