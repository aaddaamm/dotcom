#!/usr/bin/env python3
"""Generates static/adam_robinson.docx from structured content."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEAL = RGBColor(0x2A, 0x7A, 0x7A)
BLACK = RGBColor(0x0A, 0x0A, 0x0A)
MUTED = RGBColor(0x44, 0x44, 0x44)


def clear_paragraph_spacing(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def add_rule(doc):
    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2A7A7A')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_name_block(doc):
    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('adam robinson')
    run.font.name = 'JetBrains Mono'
    run.font.size = Pt(26)
    run.font.bold = False
    run.font.color.rgb = BLACK

    p2 = doc.add_paragraph()
    clear_paragraph_spacing(p2)
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run('SENIOR SOFTWARE ENGINEER')
    run2.font.name = 'JetBrains Mono'
    run2.font.size = Pt(9)
    run2.font.color.rgb = TEAL

    p3 = doc.add_paragraph()
    clear_paragraph_spacing(p3)
    p3.paragraph_format.space_after = Pt(2)
    run3 = p3.add_run('Providence, RI  ·  Open to relocation')
    run3.font.name = 'Calibri'
    run3.font.size = Pt(10)
    run3.font.color.rgb = MUTED

    p4 = doc.add_paragraph()
    clear_paragraph_spacing(p4)
    p4.paragraph_format.space_after = Pt(10)
    contact = 'adamrobinson.tech  ·  github.com/aaddaamm  ·  linkedin.com/in/adam-robinson-software  ·  adam@adamrobinson.tech'
    run4 = p4.add_run(contact)
    run4.font.name = 'Calibri'
    run4.font.size = Pt(9)
    run4.font.color.rgb = MUTED


def add_section_heading(doc, text):
    add_rule(doc)
    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.font.name = 'JetBrains Mono'
    run.font.size = Pt(8)
    run.font.color.rgb = TEAL


def add_body(doc, text, space_after=4):
    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    return p


def add_role_header(doc, company, title, period, location=None):
    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)

    r1 = p.add_run(company)
    r1.font.name = 'Calibri'
    r1.font.size = Pt(10)
    r1.font.bold = True
    r1.font.color.rgb = BLACK

    r2 = p.add_run(f'  ·  {title}')
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10)
    r2.font.bold = False
    r2.font.color.rgb = MUTED

    p2 = doc.add_paragraph()
    clear_paragraph_spacing(p2)
    p2.paragraph_format.space_after = Pt(4)
    loc_text = f'{period}' + (f'  ·  {location}' if location else '')
    r3 = p2.add_run(loc_text)
    r3.font.name = 'JetBrains Mono'
    r3.font.size = Pt(8)
    r3.font.color.rgb = MUTED


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    clear_paragraph_spacing(p)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK


def add_skill_row(doc, category, items):
    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    p.paragraph_format.space_after = Pt(3)

    r1 = p.add_run(f'{category:<12}')
    r1.font.name = 'JetBrains Mono'
    r1.font.size = Pt(9)
    r1.font.color.rgb = TEAL

    r2 = p.add_run(items)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10)
    r2.font.color.rgb = MUTED


def build():
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Default paragraph spacing
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    add_name_block(doc)

    # Summary
    add_section_heading(doc, 'Summary')
    add_body(doc,
        'Senior software engineer with over a decade of experience across fintech, healthcare, '
        'and enterprise. Full-stack, backend-leaning — specializing in Rails, Node.js, and '
        'TypeScript. Embeds with existing teams, gets up to speed fast in complex codebases, '
        'and ships reliably. Daily user of AI-assisted development tools.'
    )

    # Experience
    add_section_heading(doc, 'Experience')

    add_role_header(doc, 'MojoTech', 'Senior Software Engineer / Technical Lead', '2015 – Present', 'Providence, RI')
    add_body(doc,
        'Delivered across dozens of client engagements as an embedded senior engineer and '
        'technical lead. Selected clients:',
        space_after=6
    )

    add_role_header(doc, 'iCapital', 'Staff Augmentation / Senior Engineer', '2024 – present')
    add_body(doc,
        'Embedded on a large-scale alternative investment platform serving wealth managers. '
        'Expanded i18n support across static and database-backed content, co-designed a Rails '
        'service for bulk nominee investment processing, and led a component library migration '
        '(Supernova v1 → v2).'
    )

    add_role_header(doc, 'Healthcasts', 'Technical Lead', '2022 – 2024')
    add_body(doc,
        'Led platform modernization for a medical publishing company. Built a headless CMS '
        'publishing pipeline (Strapi + React), rebuilt AWS infrastructure, and delivered an '
        'Auth0 authentication overhaul unifying login across all platforms — unblocking a '
        'parallel AI initiative in the process.'
    )

    add_role_header(doc, 'Angi', 'Staff Augmentation / Senior Engineer', '2021 – 2022')
    add_body(doc,
        'Delivered across three separate codebases (HomeAdvisor, Handy, Angie\'s List) in a '
        'single engagement — Vue/Java, Rails/React, and Next.js/Contentful. Mentored a team '
        'of interns through their first fully shipped feature.'
    )

    add_role_header(doc, 'Shell Techworks', 'Software Engineer', '2018 – 2019')
    add_body(doc,
        'Built decommissioning tooling for end-of-life offshore oil platforms. Full-stack '
        'React/Node application. Used the Google Design Sprint process to compress scope '
        'and deliver MVP on schedule onsite in Boston.'
    )

    add_role_header(doc, 'Beacon Mutual Insurance', 'Software Engineer', '2013 – 2016', 'Warwick, RI')
    add_body(doc,
        'Built backend systems for claims and policy management in a regulated, '
        'high-reliability environment. Designed financial transaction and payment processing '
        'systems with strong correctness requirements.'
    )

    # Skills
    add_section_heading(doc, 'Skills')
    add_skill_row(doc, 'Backend', 'TypeScript, Ruby, SQL, Elixir  ·  Node.js, Ruby on Rails, Express, Phoenix')
    add_skill_row(doc, 'Frontend', 'React, SvelteKit, Vue, Next.js')
    add_skill_row(doc, 'Cloud', 'AWS  ·  Vercel  ·  GitHub Actions')
    add_skill_row(doc, 'Tools', 'Git, Prisma, Strapi, Contentful, Auth0')
    add_skill_row(doc, 'AI', 'GitHub Copilot, Claude, AMP  (daily — VS Code / Zed)')

    out = '/Users/adam/dotcom/static/adam_robinson.docx'
    doc.save(out)
    print(f'Saved: {out}')


if __name__ == '__main__':
    build()
